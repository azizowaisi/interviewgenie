"""
DOCX → structured CV JSON → LLM (via llm-service) → new DOCX for ATS optimization.
"""
from __future__ import annotations

import io
import json
import logging
import re
import threading
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

DOCX_SUFFIXES = (".docx",)
DOCX_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_SECTION_ALIASES = {
    "summary": (
        "summary",
        "professional summary",
        "profile",
        "objective",
        "about me",
        "career objective",
    ),
    "experience": (
        "experience",
        "work experience",
        "employment history",
        "professional experience",
        "employment",
        "career history",
    ),
    "skills": (
        "skills",
        "technical skills",
        "core competencies",
        "expertise",
        "technologies",
    ),
    "education": ("education", "academic", "qualifications"),
}


def _norm_heading(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip().lower().rstrip(":"))


def _match_section(line: str) -> str | None:
    h = _norm_heading(line)
    for key, aliases in _SECTION_ALIASES.items():
        for a in aliases:
            if h == a or h.startswith(a + " "):
                return key
    return None


def _collect_paragraph_lines(doc) -> list[str]:
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        lines.append(t)
    return lines


_BULLET_RE = re.compile(
    r"^[\u2022\u2023\u2043\u2219\u00B7•\-*]\s*(.+)$",
)


def _parse_experience_lines(exp_lines: list[str]) -> list[dict[str, Any]]:
    if not exp_lines:
        return []
    roles: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None

    def ensure_current() -> dict[str, Any]:
        nonlocal current
        if current is None:
            current = {"role": "", "company": "", "bullets": []}
        return current

    for line in exp_lines:
        m = _BULLET_RE.match(line)
        if m:
            cur = ensure_current()
            cur.setdefault("bullets", []).append(m.group(1).strip())
            continue

        if current and (current.get("role") or current.get("bullets")):
            roles.append(current)
            current = None

        parts = re.split(r"\s+at\s+", line, maxsplit=1, flags=re.I)
        if len(parts) == 2:
            current = {"role": parts[0].strip(), "company": parts[1].strip(), "bullets": []}
        elif "|" in line:
            a, b = line.split("|", 1)
            current = {"role": a.strip(), "company": b.strip(), "bullets": []}
        else:
            current = {"role": line.strip(), "company": "", "bullets": []}

    if current and (current.get("role") or current.get("bullets")):
        roles.append(current)
    return roles


def docx_bytes_to_structure(data: bytes) -> dict[str, Any]:
    """Extract {name, summary, experience, skills} from DOCX bytes."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = _collect_paragraph_lines(doc)
    if not lines:
        return {"name": "", "summary": "", "experience": [], "skills": []}

    name = lines[0]
    section_map: dict[str, tuple[int, int]] = {}
    i = 1
    while i < len(lines):
        sec = _match_section(lines[i])
        if sec:
            start = i + 1
            j = start
            while j < len(lines):
                nxt = _match_section(lines[j])
                if nxt and nxt != sec:
                    break
                j += 1
            section_map[sec] = (start, j)
            i = j
            continue
        i += 1

    def slice_sec(key: str) -> list[str]:
        if key not in section_map:
            return []
        a, b = section_map[key]
        return lines[a:b]

    summary_lines = slice_sec("summary")
    exp_lines = slice_sec("experience")
    skill_lines = slice_sec("skills")

    summary = " ".join(summary_lines).strip()
    experience = _parse_experience_lines(exp_lines)
    skills_flat: list[str] = []
    for sl in skill_lines:
        if "," in sl or ";" in sl:
            for part in re.split(r"[,;]", sl):
                p = part.strip()
                if p:
                    skills_flat.append(p)
        elif sl:
            skills_flat.append(sl)

    if not summary and name and len(lines) > 1:
        tail = []
        for k in ("experience", "skills", "education"):
            if k in section_map:
                a, _ = section_map[k]
                tail.append(a)
        end = min(tail) if tail else len(lines)
        if end > 1:
            summary = " ".join(lines[1:end]).strip()

    if not experience:
        rest = lines[1:]
        if rest:
            experience = _parse_experience_lines(rest[:50])

    return {
        "name": name.strip(),
        "summary": summary,
        "experience": experience,
        "skills": skills_flat,
    }


def plaintext_cv_to_structure(text: str) -> dict[str, Any]:
    """Best-effort structure when only parsed plain text exists (e.g. PDF/TXT upload)."""
    text = (text or "").strip()
    if not text:
        return {"name": "", "summary": "", "experience": [], "skills": []}
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    name = (lines[0] if lines else "")[:200]
    body = "\n".join(lines[1:]) if len(lines) > 1 else text
    body = body[:12000]
    return {
        "name": name,
        "summary": body,
        "experience": [],
        "skills": [],
    }


def validate_docx_only(filename: str | None, content_type: str | None, data: bytes) -> None:
    from fastapi import HTTPException

    fn = (filename or "").lower()
    if not fn.endswith(DOCX_SUFFIXES):
        raise HTTPException(status_code=400, detail="Only .docx files are supported (not PDF).")
    ct = (content_type or "").split(";")[0].strip().lower()
    if "pdf" in fn or (ct and "pdf" in ct):
        raise HTTPException(status_code=400, detail="PDF is not supported. Upload a .docx file only.")
    if len(data) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 15MB).")
    try:
        from docx import Document

        Document(io.BytesIO(data))
    except Exception as e:
        logger.warning("DOCX open failed: %s", e)
        raise HTTPException(status_code=400, detail="Invalid or corrupted DOCX file.") from e


def build_optimize_prompt(cv_json: dict[str, Any], job_description: str) -> str:
    cv_s = json.dumps(cv_json, ensure_ascii=False, indent=2)
    jd = (job_description or "").strip()
    return f"""You are an ATS resume optimizer.

Job description:
---
{jd}
---

Current CV (structured JSON):
---
{cv_s}
---

Tasks:
1. Improve ATS match for the job description
2. Add missing keywords naturally
3. Rewrite bullet points with strong action verbs and measurable impact where the original implies metrics
4. Keep content truthful (DO NOT invent employers, dates, degrees, roles, certifications, or projects)
5. Keep format clean, detailed, and professional (do NOT drop existing content)
6. If you are unsure about a field, keep the original content instead of removing it

Return ONLY valid JSON (no markdown fences), with this exact shape:
{{
  "name": "full name from CV or empty string",
  "summary": "2-4 sentence professional summary",
  "experience": [
    {{"role": "Job title", "company": "Company", "bullets": ["bullet1", "bullet2"]}}
  ],
  "skills": ["skill1", "skill2"],
  "suggestions": [
    "short actionable suggestion strings"
  ]
}}
"""


def parse_llm_json_object(raw: str) -> dict[str, Any]:
    from fastapi import HTTPException

    text = (raw or "").strip()
    if not text:
        raise HTTPException(status_code=502, detail="LLM did not return valid JSON.")

    def sanitize_jsonish(s: str) -> str:
        s = (s or "").strip()
        # Remove UTF-8 BOM if present.
        s = s.lstrip("\ufeff")
        # Normalize common “smart quotes” to ASCII.
        s = (
            s.replace("\u201c", '"')
            .replace("\u201d", '"')
            .replace("\u201e", '"')
            .replace("\u201f", '"')
            .replace("\u2018", "'")
            .replace("\u2019", "'")
        )
        # Remove trailing commas before } or ] (common LLM mistake).
        prev = None
        while prev != s:
            prev = s
            s = re.sub(r",\s*([}\]])", r"\1", s)
        return s

    def ensure_obj(v: Any) -> dict[str, Any] | None:
        return v if isinstance(v, dict) else None

    # 1) Best case: response is a JSON object.
    try:
        parsed = json.loads(sanitize_jsonish(text))
        obj = ensure_obj(parsed)
        if obj is not None:
            return obj
    except json.JSONDecodeError:
        pass

    # 2) Common case: JSON wrapped in a fenced block somewhere.
    for block in re.findall(r"```(?:json)?\s*([\s\S]*?)\s*```", text, flags=re.I):
        candidate = sanitize_jsonish(block)
        if not candidate:
            continue
        try:
            parsed = json.loads(candidate)
            obj = ensure_obj(parsed)
            if obj is not None:
                return obj
        except json.JSONDecodeError:
            continue

    # 3) Fallback: find the first valid JSON object substring.
    decoder = json.JSONDecoder()
    stext = sanitize_jsonish(text)
    for i, ch in enumerate(stext):
        if ch != "{":
            continue
        try:
            parsed, _end = decoder.raw_decode(stext[i:])
        except json.JSONDecodeError:
            continue
        obj = ensure_obj(parsed)
        if obj is not None:
            return obj

    logger.warning("LLM JSON parse failed; raw preview=%r", text[:800])
    raise HTTPException(status_code=502, detail="LLM did not return valid JSON.")


def normalize_llm_output(
    data: dict[str, Any],
    fallback_name: str,
    original: dict[str, Any] | None = None,
) -> dict[str, Any]:
    original = original or {}
    name = (data.get("name") or fallback_name or original.get("name") or "").strip()
    summary = (data.get("summary") or "").strip() or (original.get("summary") or "").strip()
    skills = data.get("skills") or []
    if not isinstance(skills, list):
        skills = []
    skills = [str(s).strip() for s in skills if str(s).strip()]
    if not skills:
        orig_sk = original.get("skills") or []
        if isinstance(orig_sk, list):
            skills = [str(s).strip() for s in orig_sk if str(s).strip()]
    suggestions = data.get("suggestions") or []
    if not isinstance(suggestions, list):
        suggestions = []
    suggestions = [str(s).strip() for s in suggestions if str(s).strip()]
    exp_in = data.get("experience") or []
    experience: list[dict[str, Any]] = []
    if isinstance(exp_in, list):
        for item in exp_in:
            if not isinstance(item, dict):
                continue
            role = str(item.get("role") or "").strip()
            company = str(item.get("company") or "").strip()
            bullets = item.get("bullets") or []
            if not isinstance(bullets, list):
                bullets = []
            bullets = [str(b).strip() for b in bullets if str(b).strip()]
            if role or company or bullets:
                experience.append({"role": role, "company": company, "bullets": bullets})
    if not experience:
        orig_exp = original.get("experience") or []
        if isinstance(orig_exp, list):
            for item in orig_exp:
                if isinstance(item, dict):
                    role = str(item.get("role") or "").strip()
                    company = str(item.get("company") or "").strip()
                    bullets = item.get("bullets") or []
                    if not isinstance(bullets, list):
                        bullets = []
                    bullets = [str(b).strip() for b in bullets if str(b).strip()]
                    if role or company or bullets:
                        experience.append({"role": role, "company": company, "bullets": bullets})
    return {
        "name": name,
        "summary": summary,
        "experience": experience,
        "skills": skills,
        "suggestions": suggestions,
    }


def write_structure_to_docx(path: str | Path, data: dict[str, Any]) -> None:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt

    path = Path(path)
    doc = Document()
    name = (data.get("name") or "Candidate").strip() or "Candidate"
    t = doc.add_heading(name, level=0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in t.runs:
        run.font.size = Pt(18)

    summary = (data.get("summary") or "").strip()
    if summary:
        doc.add_heading("Summary", level=1)
        p = doc.add_paragraph(summary)
        for run in p.runs:
            run.font.size = Pt(11)

    exp = data.get("experience") or []
    if exp:
        doc.add_heading("Experience", level=1)
        for job in exp:
            if not isinstance(job, dict):
                continue
            role = (job.get("role") or "").strip()
            company = (job.get("company") or "").strip()
            if role and company:
                head = f"{role} — {company}"
            else:
                head = role or company
            if head:
                hp = doc.add_paragraph()
                hr = hp.add_run(head)
                hr.bold = True
                hr.font.size = Pt(11)
            for b in job.get("bullets") or []:
                if str(b).strip():
                    bp = doc.add_paragraph(str(b).strip(), style="List Bullet")
                    for run in bp.runs:
                        run.font.size = Pt(11)

    skills = data.get("skills") or []
    if skills:
        doc.add_heading("Skills", level=1)
        sp = doc.add_paragraph(", ".join(str(s).strip() for s in skills if str(s).strip()))
        for run in sp.runs:
            run.font.size = Pt(11)

    doc.save(str(path))


# --- ephemeral download registry ---
_lock = threading.Lock()
_downloads: dict[str, tuple[str, str]] = {}
# file_id -> (user_id, absolute_path)


def register_download(user_id: str, abs_path: str, file_id: str) -> None:
    with _lock:
        if len(_downloads) > 500:
            _prune_unlocked(400)
        _downloads[file_id] = (user_id, abs_path)


def take_download_path(user_id: str, file_id: str) -> str | None:
    """Remove registry entry and return path if it belongs to user_id."""
    with _lock:
        row = _downloads.pop(file_id, None)
        if not row:
            return None
        uid, path = row
        if uid != user_id:
            _downloads[file_id] = row
            return None
        return path


def _prune_unlocked(keep: int) -> None:
    if len(_downloads) <= keep:
        return
    # drop oldest entries (dict preserves insertion order in Py3.7+)
    while len(_downloads) > keep:
        k = next(iter(_downloads))
        del _downloads[k]


async def call_llm_generate(prompt: str, llm_service_url: str, timeout: float = 120.0) -> str:
    import httpx

    base = llm_service_url.rstrip("/")
    async with httpx.AsyncClient(timeout=timeout) as client:
        resp = await client.post(f"{base}/generate", json={"prompt": prompt})
        resp.raise_for_status()
        data = resp.json()
        return (data.get("raw_answer") or data.get("text") or data.get("answer") or "").strip()
