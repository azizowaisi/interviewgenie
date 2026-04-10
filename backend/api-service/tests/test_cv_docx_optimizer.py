"""Tests for DOCX ATS optimizer helpers."""
import io

import pytest
from docx import Document
from fastapi import HTTPException

from cv_docx_optimizer import (
    docx_bytes_to_structure,
    parse_llm_json_object,
    plaintext_cv_to_structure,
    validate_docx_only,
)


def _tiny_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Ada Lovelace")
    doc.add_paragraph("Summary")
    doc.add_paragraph("Mathematician and writer.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_reject_pdf_extension():
    with pytest.raises(HTTPException) as ei:
        validate_docx_only("resume.pdf", "application/pdf", b"%PDF-1.4")
    assert ei.value.status_code == 400


def test_accept_docx():
    data = _tiny_docx()
    validate_docx_only("cv.docx", "application/octet-stream", data)


def test_structure_extracts_name():
    data = _tiny_docx()
    s = docx_bytes_to_structure(data)
    assert s["name"] == "Ada Lovelace"


def test_plaintext_structure():
    s = plaintext_cv_to_structure("Jane Doe\n\nEngineer with 5 years experience.")
    assert s["name"] == "Jane Doe"
    assert "Engineer" in s["summary"]


def test_parse_llm_json_plain_object():
    obj = parse_llm_json_object('{"name":"A","summary":"","experience":[],"skills":[],"suggestions":[]}')
    assert obj["name"] == "A"


def test_parse_llm_json_inside_fence():
    raw = "Sure — here it is:\n```json\n{\"name\":\"A\",\"summary\":\"s\",\"experience\":[],\"skills\":[],\"suggestions\":[]}\n```\n"
    obj = parse_llm_json_object(raw)
    assert obj["summary"] == "s"


def test_parse_llm_json_inside_prose_extracts_first_object():
    raw = "Some explanation.\n{\"name\":\"A\",\"summary\":\"s\",\"experience\":[],\"skills\":[],\"suggestions\":[]}\nMore text."
    obj = parse_llm_json_object(raw)
    assert obj["name"] == "A"
