from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

import main


def test_health():
    client = TestClient(main.app)
    r = client.get("/health")
    assert r.status_code == 200
    assert r.json()["status"] == "ok"


def test_render_docx_returns_docx_bytes():
    client = TestClient(main.app)
    payload = {
        "name": "Ada Lovelace",
        "summary": "Mathematician and writer.",
        "experience": [{"role": "Analyst", "company": "Example", "bullets": ["Did work"]}],
        "skills": ["Python", "Docker"],
    }
    r = client.post("/render/docx", json=payload)
    assert r.status_code == 200
    assert r.headers.get("content-type", "").startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(r.content) > 500

    doc = Document(BytesIO(r.content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    flat = "\n".join(parts)
    assert "Ada Lovelace" in flat
    assert "PROFESSIONAL SUMMARY" in flat
    assert "Python" in flat and "Docker" in flat

