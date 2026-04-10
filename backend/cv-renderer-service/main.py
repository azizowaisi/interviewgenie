from __future__ import annotations

import io
from typing import Any

from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel, Field


class ExperienceItem(BaseModel):
    role: str = ""
    company: str = ""
    bullets: list[str] = Field(default_factory=list)


class RenderCvRequest(BaseModel):
    name: str = ""
    summary: str = ""
    experience: list[ExperienceItem] = Field(default_factory=list)
    skills: list[str] = Field(default_factory=list)


app = FastAPI(title="CV Renderer Service", version="0.1.0")


@app.get("/health")
async def health() -> dict[str, str]:
    return {"status": "ok"}


def _add_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run((text or "").strip())
    r.bold = True


def _add_bullets(doc, bullets: list[str]) -> None:
    for b in bullets:
        t = (b or "").strip()
        if not t:
            continue
        doc.add_paragraph(t, style="List Bullet")


def _render_name(doc, name: str) -> None:
    if not name:
        return
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True


def _render_summary(doc, summary: str) -> None:
    if not summary:
        return
    _add_heading(doc, "Summary")
    doc.add_paragraph(summary)


def _render_experience(doc, experience: list[ExperienceItem]) -> None:
    if not experience:
        return
    _add_heading(doc, "Experience")
    for item in experience:
        parts = [(item.role or "").strip(), (item.company or "").strip()]
        header = " — ".join([x for x in parts if x])
        if header:
            p = doc.add_paragraph()
            rr = p.add_run(header)
            rr.bold = True
        _add_bullets(doc, item.bullets or [])


def _render_skills(doc, skills: list[str]) -> None:
    if not skills:
        return
    _add_heading(doc, "Skills")
    doc.add_paragraph(", ".join(skills))


def render_cv_to_docx_bytes(payload: RenderCvRequest) -> bytes:
    from docx import Document

    doc = Document()

    name = (payload.name or "").strip()
    summary = (payload.summary or "").strip()
    experience = payload.experience or []
    skills = [s.strip() for s in (payload.skills or []) if isinstance(s, str) and s.strip()]

    _render_name(doc, name)
    _render_summary(doc, summary)
    _render_experience(doc, experience)
    _render_skills(doc, skills)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.post(
    "/render/docx",
    responses={
        422: {"description": "Invalid CV JSON for rendering"},
        500: {"description": "Renderer produced empty DOCX"},
    },
)
async def render_docx(body: dict[str, Any]) -> Response:
    try:
        req = RenderCvRequest.model_validate(body)
    except Exception as e:
        raise HTTPException(status_code=422, detail="Invalid CV JSON for rendering") from e

    data = render_cv_to_docx_bytes(req)
    if not data:
        raise HTTPException(status_code=500, detail="Renderer produced empty DOCX")

    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="ats_optimized_cv.docx"'},
    )

