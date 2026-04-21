"""Polished ATS-friendly CV layout (python-docx)."""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Refined palette: deep slate blue accent, neutral body
ACCENT = RGBColor(0x1E, 0x3A, 0x5F)
ACCENT_LIGHT = RGBColor(0x3D, 0x5A, 0x80)
MUTED = RGBColor(0x64, 0x74, 0x8B)
RULE_HEX = "1E3A5F"

NAME_PT = 26
SECTION_PT = 10
BODY_PT = 11
SKILL_PT = 10
BULLET_PT = 10.5

PAGE_MARGIN_IN = 0.72
SECTION_SPACE_BEFORE = Pt(20)
SECTION_SPACE_AFTER = Pt(8)
BLOCK_AFTER = Pt(10)


def _paragraph_bottom_border(paragraph: Any, color_hex: str = RULE_HEX, sz: str = "10") -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_cell_shading(cell: Any, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _setup_page(doc: Document) -> None:
    section = doc.sections[0]
    m = Inches(PAGE_MARGIN_IN)
    section.top_margin = m
    section.bottom_margin = m
    section.left_margin = m
    section.right_margin = m

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(BODY_PT)
    normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def _add_name_block(doc: Document, name: str) -> None:
    text = (name or "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(NAME_PT)
    run.font.color.rgb = ACCENT
    run.font.name = "Calibri Light"
    _paragraph_bottom_border(p, sz="16")


def _add_headline(doc: Document, headline: str) -> None:
    text = (headline or "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)


def _add_contact_block(doc: Document, contact: list[str]) -> None:
    clean = [str(x).strip() for x in (contact or []) if str(x).strip()]
    if not clean:
        return
    # Render as a single muted line under the name. Keeps ATS parsers happy and looks polished.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    for i, part in enumerate(clean):
        if i:
            sep = p.add_run("  ·  ")
            sep.font.size = Pt(10)
            sep.font.color.rgb = MUTED
        r = p.add_run(part)
        r.font.size = Pt(10)
        r.font.color.rgb = MUTED


def _add_section_heading(doc: Document, title: str) -> None:
    t = (title or "").strip().upper()
    if not t:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = SECTION_SPACE_BEFORE
    p.paragraph_format.space_after = SECTION_SPACE_AFTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True
    run = p.add_run(t)
    run.bold = True
    run.font.size = Pt(SECTION_PT)
    run.font.color.rgb = ACCENT
    _paragraph_bottom_border(p, sz="8")


def _add_summary_body(doc: Document, summary: str) -> None:
    text = (summary or "").strip()
    if not text:
        return
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = BLOCK_AFTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.18
    for run in p.runs:
        run.font.size = Pt(BODY_PT)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def _experience_parts(item: Any) -> tuple[str, str, list[str]]:
    if isinstance(item, dict):
        role = str(item.get("role") or "").strip()
        company = str(item.get("company") or "").strip()
        bullets = item.get("bullets") or []
    else:
        role = str(getattr(item, "role", "") or "").strip()
        company = str(getattr(item, "company", "") or "").strip()
        bullets = getattr(item, "bullets", None) or []
    blist: list[str] = []
    if isinstance(bullets, list):
        blist = [str(b).strip() for b in bullets if str(b).strip()]
    return role, company, blist


def _add_experience(doc: Document, experience: list[Any]) -> None:
    if not experience:
        return
    _add_section_heading(doc, "Experience")
    for idx, item in enumerate(experience):
        role, company, bullets = _experience_parts(item)
        if not role and not company and not bullets:
            continue

        head = doc.add_paragraph()
        head.paragraph_format.space_before = Pt(14 if idx > 0 else 2)
        head.paragraph_format.space_after = Pt(2)
        head.paragraph_format.keep_with_next = True

        if role:
            rr = head.add_run(role)
            rr.bold = True
            rr.font.size = Pt(12)
            rr.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        if role and company:
            mid = head.add_run("  ·  ")
            mid.font.size = Pt(11)
            mid.font.color.rgb = MUTED
        if company:
            cc = head.add_run(company)
            cc.italic = True
            cc.font.size = Pt(11)
            cc.font.color.rgb = ACCENT_LIGHT

        for b in bullets:
            bp = doc.add_paragraph(b, style="List Bullet")
            bp.paragraph_format.space_after = Pt(3)
            bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            bp.paragraph_format.line_spacing = 1.12
            for run in bp.runs:
                run.font.size = Pt(BULLET_PT)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_skills(doc: Document, skills: list[str]) -> None:
    clean = [s.strip() for s in skills if isinstance(s, str) and s.strip()]
    if not clean:
        return
    _add_section_heading(doc, "Skills")

    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    try:
        table.allow_autofit = True
    except Exception:
        pass
    cell = table.rows[0].cells[0]
    cell.text = ""
    _set_cell_shading(cell, "F1F5F9")

    inner = cell.paragraphs[0]
    inner.paragraph_format.space_before = Pt(10)
    inner.paragraph_format.space_after = Pt(10)
    inner.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    inner.paragraph_format.line_spacing = 1.25

    # Middot-separated flow reads cleanly in ATS and on screen
    for i, skill in enumerate(clean):
        if i:
            sep = inner.add_run("  ·  ")
            sep.font.size = Pt(SKILL_PT)
            sep.font.color.rgb = MUTED
        r = inner.add_run(skill)
        r.font.size = Pt(SKILL_PT)
        r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Remove default table borders for a soft card look
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def render_cv_to_docx_bytes(
    *,
    name: str,
    headline: str,
    contact: list[str],
    summary: str,
    experience: list[Any],
    skills: list[str],
    education: list[str],
) -> bytes:
    doc = Document()
    _setup_page(doc)

    _add_name_block(doc, name)
    _add_headline(doc, headline)
    _add_contact_block(doc, contact)

    if (summary or "").strip():
        _add_section_heading(doc, "Professional summary")
        _add_summary_body(doc, summary)

    _add_experience(doc, experience)
    _add_skills(doc, skills)
    if education:
        _add_section_heading(doc, "Education")
        for e in education:
            if str(e).strip():
                p = doc.add_paragraph(str(e).strip())
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(BODY_PT)
                    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
